#These Libraries are necessary to run the pipeline.

#!pip install pandas
#!pip install xgboost
#!pip install python-docx
#!pip install textract

file_name="input_word/flower _2.docx"  # Input Document File name.
output ="output/flower2_o.docx"  # This will be the name of the final document generated . You can change this to get desired filename.   

# Importing Necessary Libraries
import pandas as pd
import textract
import pickle
import docx
import xgboost as xgb
import warnings
import json 

warnings.simplefilter(action='ignore')


text=textract.process(file_name)
text=text.decode("utf-8")

#Splitting Document by /n
lines=text.splitlines()  

df=pd.DataFrame()

#Feature Generation

df['lines']=lines
df['char_length']=df['lines'].apply(lambda x : len(x))
df['no_of_words']=df['lines'].apply(lambda x: len(x.split()))

df['index_']=df.index


def emptylines(df):
    """Calculates Number of Upcoming and preceding empty lines """
    for i in range(len(df)):
        if(i==0):
            if(df.loc[i+1]["lines"]==""):
                df.at[i,"space"]=1
            else:
                df.at[i,'space']=0
            i+=1
            continue

        if(i==len(df)-1):
            if(df.loc[i-1]["lines"]==""):
                df.at[i,"space"]=1
            else:
                df.at[i,'space']=0
            i+=1
            continue

        if((df.loc[i-1]['lines']=="") ^(df.loc[i+1]['lines']=="")):
           df.at[i,'space']=1
        elif((df.loc[i-1]['lines']=="") and (df.loc[i+1]['lines']=="")):
           df.at[i,'space']=2
        else:
           df.at[i,'space']=0
        i+=1

def caps(text):
    """ Calculates the percentage of Capital letters in word."""
    caps_count=0
    for i in text:
        if(i.isupper()):
            caps_count+=1
    if caps_count==0 or len(text)==0:
        return 0
    else:
        return  caps_count/len(text)
    
def feature_generator(df):
    """
    This Function generates essential features for the model . 
    They are number of characters and words in previous line ,next line ,previous of previous line ,next to next line and the number of capital letters in the next sentence. 
    """
    len_df=len(df)
    for i in range(len(df)):
        if(i==0):
            df.at[i,"next_char"]=df.loc[i+1]["char_length"]
            df.at[i,"next_words"]=df.loc[i+1]["no_of_words"]
            df.at[i,"next_next_char"]=df.loc[i+2]["char_length"]
            df.at[i,"next_next_words"]=df.loc[i+2]["no_of_words"]
            df.at[i,"next_caps"]=caps(df.at[i+1,"lines"])
            df.at[i,"prev_char"]=0
            df.at[i,"prev_words"]=0
            df.at[i,"prev_prev_char"]=0
            df.at[i,"prev_prev_words"]=0
            i+=1
            continue
            
        if(i==1):
            df.at[i,"next_char"]=df.loc[i+1]["char_length"]
            df.at[i,"next_words"]=df.loc[i+1]["no_of_words"]
            df.at[i,"next_next_char"]=df.loc[i+2]["char_length"]
            df.at[i,"next_next_words"]=df.loc[i+2]["no_of_words"]
            df.at[i,"prev_char"]=df.loc[i-1]["char_length"]
            df.at[i,"prev_words"]=df.loc[i-1]["no_of_words"]
            df.at[i,"prev_prev_char"]=0
            df.at[i,"prev_prev_words"]=0      
            
            df.at[i,"next_caps"]=caps(df.at[i+1,"lines"])
            i+=1
            continue
            
        if(i==len_df-2):
            df.at[i,"next_char"]=df.loc[i+1]["char_length"]
            df.at[i,"next_words"]=df.loc[i+1]["no_of_words"]
            df.at[i,"next_next_char"]=0
            df.at[i,"next_next_words"]=0
            df.at[i,"prev_char"]=df.loc[i-1]["char_length"]
            df.at[i,"prev_words"]=df.loc[i-1]["no_of_words"]
            df.at[i,"prev_prev_char"]=df.loc[i-2]["char_length"]
            df.at[i,"prev_prev_words"]=df.loc[i-2]["no_of_words"]
            
            df.at[i,"next_caps"]=caps(df.at[i+1,"lines"])
            i+=1
            continue
            
            
        if(i==len_df-1):
            df.at[i,"next_char"]=0
            df.at[i,"next_words"]=0
            df.at[i,"next_next_char"]=0
            df.at[i,"next_next_words"]=0
            df.at[i,"prev_char"]=df.loc[i-1]["char_length"]
            df.at[i,"prev_words"]=df.loc[i-1]["no_of_words"]
            df.at[i,"prev_prev_char"]=df.loc[i-2]["char_length"]
            df.at[i,"prev_prev_words"]=df.loc[i-2]["no_of_words"]
            
            df.at[i,"next_caps"]=0
            i+=1
            continue
        
        df.at[i,"next_char"]=df.loc[i+1]["char_length"]
        df.at[i,"next_words"]=df.loc[i+1]["no_of_words"]
        df.at[i,"next_next_char"]=df.loc[i+2]["char_length"]
        df.at[i,"next_next_words"]=df.loc[i+2]["no_of_words"]
        df.at[i,"prev_char"]=df.loc[i-1]["char_length"]
        df.at[i,"prev_words"]=df.loc[i-1]["no_of_words"]
        df.at[i,"prev_prev_char"]=df.loc[i-2]["char_length"]
        df.at[i,"prev_prev_words"]=df.loc[i-2]["no_of_words"]
        
        df.at[i,"next_caps"]=caps(df.at[i+1,"lines"])
        i+=1
            

#Applying the functions created 
df['caps']=df['lines'].apply(caps)
emptylines(df)
feature_generator(df)


"""Loading the XGBOOST Model trained using the data generated with the above method of generating features.
    6 Wikipedia articles were handpicked and were annotated for this purpose."""

with open('model_pkl' , 'rb') as f:
    model = pickle.load(f)


#Splitting the Dataset into non-empty lines and empty lines. empty lines are preserved to add to the final document.
to_send=df[df["lines"]!=""]

not_send=df[df["lines"]==""]

used=['char_length', 'no_of_words',
       'caps', 'space', 'next_char', 'next_words', 'next_next_char',
       'next_next_words', 'next_caps', 'prev_char', 'prev_words',
       'prev_prev_char', 'prev_prev_words']


predictions=model.predict(to_send[used])  # Prediction by the XGBOOST Model 


to_send["label"]=predictions
not_send['label']=0

# Concatanating the dataset.
text=pd.concat([to_send,not_send])
text.sort_values(by="index_",inplace=True)



my_doc = docx.Document()  #Creating a word document.

before=True
for i in range(len(text)):
    word=text.loc[i,"lines"]
    label=text.loc[i,"label"]
    if(word!=""):
        if(label==1): # If Label is Heading do this
            if(i!=0):
                my_doc.add_paragraph("<<END SECTION>>")
            my_doc.add_heading("<<START HEADING>> "+word+" <<END HEADING>>",2)
            before=True
        else:
            if(before==True): #If Label is section do this
                my_doc.add_paragraph("<<START SECTION>>")
            my_doc.add_paragraph(word)
            before=False
            if(i+1 <len(text)):
                if(text.loc[i+1,"label"]==1):
                    pass
        
    else: #For Empty Lines
        #my_doc.add_paragraph(word)
        pass

        
#Saving the document.
my_doc.save(output)